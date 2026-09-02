"""DOCX parser extracting text, paragraph styles, section headings, and table contents."""

from __future__ import annotations

import io
import logging
from typing import List, Tuple

from modules.rag.src.models import RawSection

logger = logging.getLogger(__name__)


def parse_docx(file_bytes: bytes) -> Tuple[List[RawSection], List[str]]:
    """Parse a DOCX document into raw sections partitioned by Heading styles.

    Args:
        file_bytes: Raw bytes of the uploaded DOCX document.

    Returns:
        Tuple of (raw_sections, detected_chapters)
    """
    sections: List[RawSection] = []
    chapters: List[str] = []

    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        
        current_heading: str | None = None
        current_paras: List[str] = []
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name if p.style else ""
            
            # Check if paragraph is a Heading (e.g. "Heading 1", "Heading 2", "Title")
            if style_name.lower().startswith("heading") or style_name.lower() == "title":
                # Flush previous accumulated section
                if current_paras:
                    sections.append(
                        RawSection(
                            section_title=current_heading,
                            page_or_slide=None,
                            raw_text="\n\n".join(current_paras),
                            metadata={"style": style_name}
                        )
                    )
                    current_paras = []
                
                current_heading = text
                if style_name.lower() in ("heading 1", "title") and text not in chapters:
                    chapters.append(text)
            else:
                current_paras.append(text)

        # Also extract table text into raw sections
        for t in doc.tables:
            table_rows = []
            for row in t.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_rows.append(row_text)
            if table_rows:
                current_paras.append("\n".join(table_rows))

        # Flush final section
        if current_paras:
            sections.append(
                RawSection(
                    section_title=current_heading,
                    page_or_slide=None,
                    raw_text="\n\n".join(current_paras),
                    metadata={}
                )
            )

    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")

    return sections, chapters
